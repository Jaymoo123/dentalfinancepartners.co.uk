#!/usr/bin/env python
"""Build the signing-copy .docx from its Markdown source using python-docx.

Reproduces the contract as a polished Word document (US Letter, tight margins,
styled Heading 1/2 so collapsible headings work in Word and in Google Docs
"Pageless" view, a centred title block, shaded table headers, blockquote
styling and a page-number footer). No pandoc required.

Usage:
    python build_signing_docx.py [input.md] [output.docx]
Defaults to the FOR_SIGNATURE pair in this folder.
"""
import re
import sys

from docx import Document
from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

DEFAULT_IN = "Lead_Generation_and_Data_Sharing_Agreement_FOR_SIGNATURE.md"
DEFAULT_OUT = "Lead_Generation_and_Data_Sharing_Agreement_FOR_SIGNATURE.docx"

INLINE_RE = re.compile(r"(\*\*.+?\*\*|\*.+?\*)")
SEP_CELL_RE = re.compile(r"^:?-{2,}:?$")

INK = RGBColor(0x1F, 0x33, 0x55)        # dark navy for headings
ACCENT = "8AA0C8"                        # heading divider rule
HEADER_FILL = "1F3355"                   # table header row fill (dark)
HEADER_TEXT = RGBColor(0xFF, 0xFF, 0xFF) # table header text
BORDER = "C9CFDA"                        # light table borders
QUOTE_BAR = "8AA0C8"                     # blockquote left accent
GREY = RGBColor(0x6B, 0x72, 0x80)


# ---- low-level OOXML helpers ------------------------------------------------

def _shade(el, fill):
    pr = el.get_or_add_tcPr() if el.tag.endswith("}tc") else el
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pr.append(shd)


def _cell_shade(cell, fill):
    _shade(cell._tc, fill)


def _table_borders(table, color=BORDER, sz=4):
    tblPr = table._tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), str(sz))
        e.set(qn("w:space"), "0")
        e.set(qn("w:color"), color)
        borders.append(e)
    tblPr.append(borders)


def _mark_header_row(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)


def _bottom_border(paragraph, color=ACCENT, sz=6, space=6):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), str(sz))
    bottom.set(qn("w:space"), str(space))
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)


def _left_bar(paragraph, color=QUOTE_BAR, sz=18, space=10):
    pPr = paragraph._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    left = OxmlElement("w:left")
    left.set(qn("w:val"), "single")
    left.set(qn("w:sz"), str(sz))
    left.set(qn("w:space"), str(space))
    left.set(qn("w:color"), color)
    pBdr.append(left)
    pPr.append(pBdr)


def _page_field(paragraph, instr):
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr_el = OxmlElement("w:instrText")
    instr_el.set(qn("xml:space"), "preserve")
    instr_el.text = instr
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.append(begin)
    run._r.append(instr_el)
    run._r.append(end)


def _footer(section, title):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run(f"{title}    |    Page ")
    _page_field(p, "PAGE")
    p.add_run(" of ")
    _page_field(p, "NUMPAGES")
    for r in p.runs:
        r.font.size = Pt(8.5)
        r.font.name = "Calibri"
        r.font.color.rgb = GREY


# ---- markdown helpers -------------------------------------------------------

def add_inline(paragraph, text, base_bold=False, base_italic=False, color=None):
    for token in INLINE_RE.split(text):
        if not token:
            continue
        bold, italic, body = base_bold, base_italic, token
        if token.startswith("**") and token.endswith("**"):
            bold, body = True, token[2:-2]
        elif token.startswith("*") and token.endswith("*"):
            italic, body = True, token[1:-1]
        run = paragraph.add_run(body)
        run.bold = bold or None
        run.italic = italic or None
        if color is not None:
            run.font.color.rgb = color


def split_row(line):
    cells = line.strip().split("|")
    if cells and cells[0] == "":
        cells = cells[1:]
    if cells and cells[-1] == "":
        cells = cells[:-1]
    return [c.strip() for c in cells]


def is_separator_row(line):
    cells = split_row(line)
    return bool(cells) and all(SEP_CELL_RE.match(c or "-") for c in cells)


def flush_table(doc, rows):
    if not rows:
        return
    header = split_row(rows[0])
    body = [split_row(r) for r in rows[1:]]
    ncol = max([len(header)] + [len(r) for r in body])
    table = doc.add_table(rows=0, cols=ncol)
    table.style = "Table Grid"
    _table_borders(table)

    def fill(cells, header_row):
        tr = table.add_row()
        if header_row:
            _mark_header_row(tr)
        for i in range(ncol):
            txt = cells[i] if i < len(cells) else ""
            cell = tr.cells[i]
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            para.paragraph_format.space_before = Pt(2)
            if header_row:
                _cell_shade(cell, HEADER_FILL)
                add_inline(para, txt, base_bold=True, color=HEADER_TEXT)
            else:
                add_inline(para, txt)

    fill(header, True)
    for r in body:
        fill(r, False)


# ---- main -------------------------------------------------------------------

def main():
    src = sys.argv[1] if len(sys.argv) > 1 else DEFAULT_IN
    out = sys.argv[2] if len(sys.argv) > 2 else DEFAULT_OUT

    with open(src, encoding="utf-8") as fh:
        lines = fh.read().splitlines()

    doc = Document()

    sec = doc.sections[0]
    sec.page_width, sec.page_height = Inches(8.5), Inches(11)
    sec.left_margin = sec.right_margin = Inches(0.6)
    sec.top_margin = sec.bottom_margin = Inches(0.6)
    sec.footer_distance = Inches(0.35)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.12

    h1 = doc.styles["Heading 1"]
    h1.font.name = "Calibri"
    h1.font.size = Pt(15)
    h1.font.bold = True
    h1.font.color.rgb = INK
    h1.paragraph_format.space_before = Pt(14)
    h1.paragraph_format.space_after = Pt(6)
    h1.paragraph_format.keep_with_next = True

    h2 = doc.styles["Heading 2"]
    h2.font.name = "Calibri"
    h2.font.size = Pt(12)
    h2.font.bold = True
    h2.font.color.rgb = INK
    h2.paragraph_format.space_before = Pt(10)
    h2.paragraph_format.space_after = Pt(3)
    h2.paragraph_format.keep_with_next = True

    _footer(sec, "Lead Generation and Data Sharing Agreement")

    table_buf = []
    title_done = False
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()

        if line.lstrip().startswith("|") and "|" in line.lstrip()[1:]:
            table_buf.append(line)
            i += 1
            continue
        if table_buf:
            flush_table(doc, [r for r in table_buf if not is_separator_row(r)])
            table_buf = []

        stripped = line.strip()
        if not stripped or stripped == "---":
            i += 1
            continue

        if stripped.startswith("# "):
            text = stripped[2:].strip()
            if not title_done:
                # Document title block: centred, larger, ruled.
                h = doc.add_heading(level=1)
                h.alignment = WD_ALIGN_PARAGRAPH.CENTER
                add_inline(h, text)
                for r in h.runs:
                    r.font.size = Pt(20)
                _bottom_border(h, sz=8, space=8)
                title_done = True
            else:
                doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
                h = doc.add_heading(level=1)
                add_inline(h, text)
                _bottom_border(h)
        elif stripped.startswith("## "):
            h = doc.add_heading(level=2)
            add_inline(h, stripped[3:].strip())
        elif stripped.startswith(">"):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.3)
            p.paragraph_format.space_before = Pt(4)
            p.paragraph_format.space_after = Pt(4)
            _left_bar(p)
            add_inline(p, stripped.lstrip("> ").strip(), base_italic=True)
        else:
            p = doc.add_paragraph()
            if stripped.startswith("Dated "):
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            add_inline(p, stripped)
        i += 1

    if table_buf:
        flush_table(doc, [r for r in table_buf if not is_separator_row(r)])

    doc.save(out)
    print(f"Wrote {out} ({len(doc.paragraphs)} paragraphs, {len(doc.tables)} tables)")


if __name__ == "__main__":
    main()
